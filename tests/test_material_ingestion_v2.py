from pathlib import Path

import pytest
from amadeus.core.material_ingestion import (
    ingest_image,
    ingest_link,
    ingest_material,
)


def test_pdf_extraction(tmp_path):
    try:
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf not installed")

    pdf_path = tmp_path / "test.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    # pypdf blank pages have no text, so we test the flow
    with open(pdf_path, "wb") as f:
        writer.write(f)

    result = ingest_material(pdf_path, tmp_path / "context")
    assert result.status in ("ingested", "partial")
    assert result.page_count == 1
    assert result.context_path is not None


def test_docx_extraction(tmp_path):
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    docx_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("Another paragraph with details.")
    doc.save(str(docx_path))

    result = ingest_material(docx_path, tmp_path / "context")
    assert result.status == "ingested"
    assert result.context_path is not None
    content = Path(result.context_path).read_text(encoding="utf-8")
    assert "Test Heading" in content
    assert "test paragraph" in content


def test_docx_with_table(tmp_path):
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    docx_path = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("Intro text")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header1"
    table.cell(0, 1).text = "Header2"
    table.cell(1, 0).text = "Val1"
    table.cell(1, 1).text = "Val2"
    doc.save(str(docx_path))

    result = ingest_material(docx_path, tmp_path / "context")
    assert result.status == "ingested"
    assert any("table" in n.lower() for n in result.extraction_notes)


def test_image_ingestion(tmp_path):
    img_path = tmp_path / "screenshot.png"
    # Create a minimal valid PNG (1x1 pixel)
    import struct
    import zlib

    def _minimal_png():
        sig = b"\x89PNG\r\n\x1a\n"
        ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        ihdr_crc = zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF
        ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", ihdr_crc)
        raw = b"\x00\x00\x00\x00"
        idat_data = zlib.compress(raw)
        idat_crc = zlib.crc32(b"IDAT" + idat_data) & 0xFFFFFFFF
        idat = struct.pack(">I", len(idat_data)) + b"IDAT" + idat_data + struct.pack(">I", idat_crc)
        iend_crc = zlib.crc32(b"IEND") & 0xFFFFFFFF
        iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", iend_crc)
        return sig + ihdr + idat + iend

    img_path.write_bytes(_minimal_png())

    result = ingest_image(img_path, tmp_path / "context")
    assert result.status == "partial"
    assert result.context_path is not None
    assert result.extraction_confidence < 1.0
    content = Path(result.context_path).read_text(encoding="utf-8")
    assert "Image Material" in content


def test_image_unsupported_format(tmp_path):
    f = tmp_path / "file.xyz"
    f.write_bytes(b"not an image")
    result = ingest_image(f, tmp_path / "context")
    assert result.status == "unsupported_type"


def test_link_ingestion(tmp_path):
    result = ingest_link(
        "https://example.com/docs",
        tmp_path / "context",
        purpose="API documentation",
    )
    assert result.status == "partial"
    assert result.context_path is not None
    content = Path(result.context_path).read_text(encoding="utf-8")
    assert "example.com" in content
    assert "API documentation" in content


def test_text_still_works(tmp_path):
    f = tmp_path / "notes.txt"
    f.write_text("Build an API with auth", encoding="utf-8")
    result = ingest_material(f, tmp_path / "context")
    assert result.status == "ingested"
    assert result.extraction_confidence == 1.0


def test_missing_source(tmp_path):
    result = ingest_material(tmp_path / "nonexistent.txt", tmp_path / "context")
    assert result.status == "missing_source"
