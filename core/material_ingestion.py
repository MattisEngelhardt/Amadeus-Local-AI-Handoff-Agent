from __future__ import annotations

import hashlib
import importlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

MaterialIngestionStatus = Literal[
    "ingested",
    "partial",
    "dependency_missing",
    "adapter_stub",
    "unsupported_type",
    "missing_source",
    "failed",
]

TEXT_SUFFIXES = {".md", ".markdown", ".txt"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff", ".tif"}


class MaterialIngestionError(RuntimeError):
    """Base error for expected material ingestion failures."""


class MissingMaterialDependency(MaterialIngestionError):
    """Raised when an optional adapter dependency is not installed."""


class MaterialAdapterStub(MaterialIngestionError):
    """Raised when an adapter is intentionally scaffolded but not implemented."""


class UnsupportedMaterialType(MaterialIngestionError):
    """Raised when no adapter exists for the source material type."""


@dataclass(frozen=True)
class MaterialIngestionResult:
    source_id: str
    original_path: str
    context_path: str | None
    status: MaterialIngestionStatus
    extraction_notes: tuple[str, ...] = field(default_factory=tuple)
    extraction_confidence: float = 1.0
    page_count: int | None = None

    def as_dict(self) -> dict[str, object]:
        return {
            "source_id": self.source_id,
            "original_path": self.original_path,
            "context_path": self.context_path,
            "status": self.status,
            "extraction_notes": list(self.extraction_notes),
            "extraction_confidence": self.extraction_confidence,
            "page_count": self.page_count,
        }


def ingest_material(
    source_path: str | Path,
    context_dir: str | Path,
    source_id: str | None = None,
) -> MaterialIngestionResult:
    source = Path(source_path)
    resolved_original = _safe_resolve(source)
    normalized_source_id = normalize_source_id(source_id or _build_source_id(source))

    if not source.exists():
        return MaterialIngestionResult(
            source_id=normalized_source_id,
            original_path=str(resolved_original),
            context_path=None,
            status="missing_source",
            extraction_notes=("Source material does not exist.",),
        )

    if not source.is_file():
        return MaterialIngestionResult(
            source_id=normalized_source_id,
            original_path=str(resolved_original),
            context_path=None,
            status="failed",
            extraction_notes=("Source material is not a file.",),
        )

    try:
        extracted_text, notes, confidence, page_count = _extract_text(source)
    except MissingMaterialDependency as exc:
        return _failure_result(
            normalized_source_id,
            resolved_original,
            "dependency_missing",
            str(exc),
        )
    except MaterialAdapterStub as exc:
        return _failure_result(normalized_source_id, resolved_original, "adapter_stub", str(exc))
    except UnsupportedMaterialType as exc:
        return _failure_result(
            normalized_source_id,
            resolved_original,
            "unsupported_type",
            str(exc),
        )
    except UnicodeDecodeError as exc:
        return _failure_result(
            normalized_source_id,
            resolved_original,
            "failed",
            f"Text material must be UTF-8 encoded: {exc}",
        )
    except MaterialIngestionError as exc:
        return _failure_result(
            normalized_source_id,
            resolved_original,
            "failed",
            str(exc),
        )

    context_root = Path(context_dir)
    context_root.mkdir(parents=True, exist_ok=True)
    context_path = context_root / f"{normalized_source_id}.md"
    context_path.write_text(
        _format_context_document(
            source_id=normalized_source_id,
            original_path=resolved_original,
            extracted_text=extracted_text,
            page_count=page_count,
        ),
        encoding="utf-8",
    )

    status: MaterialIngestionStatus = "ingested"
    if confidence < 0.8:
        status = "partial"

    return MaterialIngestionResult(
        source_id=normalized_source_id,
        original_path=str(resolved_original),
        context_path=str(context_path),
        status=status,
        extraction_notes=tuple(notes),
        extraction_confidence=confidence,
        page_count=page_count,
    )


def ingest_image(
    source_path: str | Path,
    context_dir: str | Path,
    source_id: str | None = None,
) -> MaterialIngestionResult:
    source = Path(source_path)
    resolved_original = _safe_resolve(source)
    normalized_source_id = normalize_source_id(source_id or _build_source_id(source))

    if not source.exists() or not source.is_file():
        return _failure_result(
            normalized_source_id,
            resolved_original,
            "missing_source",
            "Image source does not exist or is not a file.",
        )

    suffix = source.suffix.lower()
    if suffix not in IMAGE_SUFFIXES:
        return _failure_result(
            normalized_source_id,
            resolved_original,
            "unsupported_type",
            f"Unsupported image format '{suffix}'.",
        )

    notes = []
    width, height = 0, 0
    file_size = source.stat().st_size

    try:
        from PIL import Image

        with Image.open(source) as img:
            width, height = img.size
            fmt = img.format or suffix.lstrip(".")
            notes.append(f"Image: {width}x{height}, format={fmt}")
    except ImportError:
        notes.append("Pillow not installed; image dimensions unavailable.")
    except Exception as exc:
        notes.append(f"Could not read image metadata: {exc}")

    notes.append(f"File size: {file_size} bytes")
    notes.append("Image material requires manual review for visual content description.")

    context_root = Path(context_dir)
    context_root.mkdir(parents=True, exist_ok=True)
    context_path = context_root / f"{normalized_source_id}.md"

    content = (
        f"# Image Material: {normalized_source_id}\n\n"
        f"- Source ID: `{normalized_source_id}`\n"
        f"- Original path: `{resolved_original}`\n"
        f"- Format: {suffix.lstrip('.')}\n"
        f"- Dimensions: {width}x{height}\n"
        f"- File size: {file_size} bytes\n\n"
        "## Visual Content\n\n"
        "_This image has not been analyzed for visual content. "
        "Manual review or local vision model required._\n\n"
        "## Notes\n\n"
    )
    for note in notes:
        content += f"- {note}\n"

    context_path.write_text(content, encoding="utf-8")

    return MaterialIngestionResult(
        source_id=normalized_source_id,
        original_path=str(resolved_original),
        context_path=str(context_path),
        status="partial",
        extraction_notes=tuple(notes),
        extraction_confidence=0.3,
    )


def ingest_link(
    url: str,
    context_dir: str | Path,
    purpose: str = "",
    source_id: str | None = None,
) -> MaterialIngestionResult:
    normalized_source_id = normalize_source_id(
        source_id or f"link-{hashlib.sha256(url.encode()).hexdigest()[:12]}"
    )

    context_root = Path(context_dir)
    context_root.mkdir(parents=True, exist_ok=True)
    context_path = context_root / f"{normalized_source_id}.md"

    content = (
        f"# Link Material: {normalized_source_id}\n\n"
        f"- Source ID: `{normalized_source_id}`\n"
        f"- URL: `{url}`\n"
        f"- Purpose: {purpose or 'Not specified'}\n\n"
        "## Content\n\n"
        "_Link content has not been fetched. "
        "Fetching requires explicit user approval and is only done for supported, safe URLs._\n"
    )

    context_path.write_text(content, encoding="utf-8")

    notes = [f"Link registered: {url}"]
    if purpose:
        notes.append(f"Purpose: {purpose}")
    notes.append("Content not fetched; requires explicit fetch approval.")

    return MaterialIngestionResult(
        source_id=normalized_source_id,
        original_path=url,
        context_path=str(context_path),
        status="partial",
        extraction_notes=tuple(notes),
        extraction_confidence=0.1,
    )


def normalize_source_id(source_id: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", source_id.strip().lower()).strip("-")
    slug = re.sub(r"-{2,}", "-", slug)
    return slug[:96] or "material"


def _build_source_id(source_path: Path) -> str:
    stem = normalize_source_id(source_path.stem)
    if source_path.exists() and source_path.is_file():
        digest = hashlib.sha256(source_path.read_bytes()).hexdigest()[:12]
    else:
        digest = hashlib.sha256(str(source_path).encode("utf-8")).hexdigest()[:12]
    return f"{stem}-{digest}"


def _extract_text(source: Path) -> tuple[str, list[str], float, int | None]:
    suffix = source.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        text, notes = _extract_plain_text(source)
        return text, list(notes), 1.0, None
    if suffix in PDF_SUFFIXES:
        return _extract_pdf_text(source)
    if suffix in DOCX_SUFFIXES:
        return _extract_docx_text(source)
    if suffix in IMAGE_SUFFIXES:
        raise UnsupportedMaterialType(
            f"Image files should use ingest_image(), not ingest_material(). "
            f"Use the image adapter for '{suffix}' files."
        )

    raise UnsupportedMaterialType(
        f"Unsupported material type '{suffix or '<none>'}'. "
        "Supported: .md, .markdown, .txt, .pdf, .docx. Images use ingest_image()."
    )


def _extract_plain_text(source: Path) -> tuple[str, tuple[str, ...]]:
    raw_text = source.read_text(encoding="utf-8")
    text = _normalize_extracted_text(raw_text)
    notes = ["Plain text material ingested as UTF-8."]
    if not text:
        notes.append("No non-whitespace text was extracted.")
    return text, tuple(notes)


def _extract_pdf_text(source: Path) -> tuple[str, list[str], float, int | None]:
    _require_optional_dependency(
        module_name="pypdf",
        package_name="pypdf",
        material_type="PDF",
    )

    from pypdf import PdfReader

    notes: list[str] = []
    try:
        reader = PdfReader(str(source))
    except Exception as exc:
        raise MaterialIngestionError(f"Could not read PDF: {exc}") from exc

    page_count = len(reader.pages)
    notes.append(f"PDF with {page_count} pages.")

    text_parts: list[str] = []
    pages_with_text = 0
    for i, page in enumerate(reader.pages, 1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
            notes.append(f"Page {i}: text extraction failed.")

        if page_text.strip():
            pages_with_text += 1
            text_parts.append(f"--- Page {i} ---\n{page_text.strip()}")
        else:
            notes.append(f"Page {i}: no extractable text (may be scanned/image).")

    full_text = _normalize_extracted_text("\n\n".join(text_parts))

    if pages_with_text == 0:
        notes.append("No text extracted from any page. Document may be scanned/image-only.")
        confidence = 0.1
    elif pages_with_text < page_count:
        confidence = round(pages_with_text / page_count, 2)
        notes.append(f"Text extracted from {pages_with_text}/{page_count} pages.")
    else:
        confidence = 0.95
        notes.append("Text extracted from all pages.")

    return full_text, notes, confidence, page_count


def _extract_docx_text(source: Path) -> tuple[str, list[str], float, int | None]:
    _require_optional_dependency(
        module_name="docx",
        package_name="python-docx",
        material_type="DOCX",
    )

    from docx import Document

    notes: list[str] = []
    try:
        doc = Document(str(source))
    except Exception as exc:
        raise MaterialIngestionError(f"Could not read DOCX: {exc}") from exc

    sections: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            level = style_name.replace("Heading", "").strip()
            try:
                heading_level = int(level)
            except ValueError:
                heading_level = 1
            sections.append(f"{'#' * heading_level} {text}")
        else:
            sections.append(text)

    table_count = len(doc.tables)
    if table_count > 0:
        notes.append(f"Document contains {table_count} table(s).")
        for i, table in enumerate(doc.tables, 1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                sections.append(f"\n### Table {i}\n")
                header = rows[0]
                sections.append(f"| {header} |")
                sections.append(f"| {' | '.join(['---'] * len(rows[0].split(' | ')))} |")
                for row_text in rows[1:]:
                    sections.append(f"| {row_text} |")

    core_props = doc.core_properties
    meta_parts = []
    if core_props.title:
        meta_parts.append(f"Title: {core_props.title}")
    if core_props.author:
        meta_parts.append(f"Author: {core_props.author}")
    if core_props.subject:
        meta_parts.append(f"Subject: {core_props.subject}")
    if meta_parts:
        notes.append(f"Metadata: {'; '.join(meta_parts)}")

    full_text = _normalize_extracted_text("\n\n".join(sections))
    para_count = len(doc.paragraphs)
    notes.append(f"Extracted {para_count} paragraphs and {table_count} tables.")

    confidence = 0.95 if full_text else 0.3

    return full_text, notes, confidence, None


def _require_optional_dependency(module_name: str, package_name: str, material_type: str) -> None:
    try:
        importlib.import_module(module_name)
    except ImportError as exc:
        raise MissingMaterialDependency(
            f"{material_type} ingestion requires optional dependency '{package_name}'. "
            "Install it before enabling this adapter."
        ) from exc


def _normalize_extracted_text(raw_text: str) -> str:
    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    normalized = "\n".join(lines).strip()
    return re.sub(r"\n{3,}", "\n\n", normalized)


def _format_context_document(
    source_id: str,
    original_path: Path,
    extracted_text: str,
    page_count: int | None = None,
) -> str:
    body = extracted_text or "_No extractable text content._"
    meta = (
        f"# Material Context: {source_id}\n\n"
        f"- Source ID: `{source_id}`\n"
        f"- Original path: `{original_path}`\n"
    )
    if page_count is not None:
        meta += f"- Pages: {page_count}\n"
    meta += f"\n## Extracted Text\n\n{body}\n"
    return meta


def _failure_result(
    source_id: str,
    original_path: Path,
    status: MaterialIngestionStatus,
    note: str,
) -> MaterialIngestionResult:
    return MaterialIngestionResult(
        source_id=source_id,
        original_path=str(original_path),
        context_path=None,
        status=status,
        extraction_notes=(note,),
    )


def _safe_resolve(path: Path) -> Path:
    try:
        return path.resolve()
    except OSError:
        return path.absolute()
